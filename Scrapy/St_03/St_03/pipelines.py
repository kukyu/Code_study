# Define your item pipelines here
#
# Don't forget to add your pipeline to the ITEM_PIPELINES setting
# See: https://docs.scrapy.org/en/latest/topics/item-pipeline.html


# useful for handling different item types with a single interface
from itemadapter import ItemAdapter
from docx import Document

class St03Pipeline:
    doc = Document()
    def process_item(self, item, spider):
        chapterName = item["chapterName"]
        content = item["content"]
        self.doc.add_heading(chapterName, level=1)
        self.doc.add_paragraph(content)
        self.doc.save("./同学两亿岁.docx")

        return item

class St03Pipeline2:
    def open_spider(self, spider):
        self.file = open("./同学两亿岁.txt", "w",encoding="utf-8")
    def process_item(self, item, spider):
        chapterName = item["chapterName"]
        content = item["content"]
        self.file.write(chapterName + "\n" + content + "\n\n")
        return item
    def close_spider(self, spider):
        self.file.close()